"""
services/permit/s1_generate.py — PTW document generation helpers.

Extracted from S1.py:
  - build_doc_data()          — converts form_data booleans to tick marks
  - generate_ptw_pdf()        — primary PDF generation (ReportLab overlay + DOCX fallback)
  - generate_ptw_docx_from_template() — DOCX template fill with green-bold formatting

All logic is verbatim from S1.py.
Only change: Streamlit imports removed; supabase_link → app.database.
"""
from __future__ import annotations

import logging
import os
import re
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Optional

from app.database import get_supabase_client

logger = logging.getLogger(__name__)

# Template constants (bucket: ptw-templates, file: PDF_Electrical_PTW_TEMPLATE.pdf)
SUPABASE_TEMPLATE_BUCKET = "ptw-templates"
TEMPLATE_PDF_FILE_NAME = "PDF_Electrical_PTW_TEMPLATE.pdf"
TEMPLATE_FILE_NAME = "Electrical_PTW_TEMPLATE.docx"

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")


# ---------------------------------------------------------------------------
# Tick mark helper (verbatim from S1.py)
# ---------------------------------------------------------------------------

def _tick(val: Any) -> str:
    if val is None:
        return ""
    if isinstance(val, bool):
        return "\u2713" if val else ""
    if isinstance(val, str):
        return "\u2713" if val.upper() in ("Y", "YES", "TRUE", "1") else ""
    return "\u2713" if bool(val) else ""


# ---------------------------------------------------------------------------
# build_doc_data (verbatim from S1.py)
# ---------------------------------------------------------------------------

def build_doc_data(form_data: dict) -> dict:
    """
    Convert form_data to document-ready data.
    - Checkbox booleans (hz_, rk_, ppe_, sp_, ap_ prefixes) → tick marks
    - undertaking_accept → tick mark
    - All other values → stringified
    """
    doc_data: dict = {}
    for k, v in form_data.items():
        if k.startswith(("hz_", "rk_", "ppe_", "sp_", "ap_")) and not k.endswith("_text"):
            doc_data[k] = _tick(v)
        elif k == "undertaking_accept":
            doc_data[k] = _tick(v)
        else:
            doc_data[k] = "" if v is None else str(v)
    return doc_data


# ---------------------------------------------------------------------------
# Template download
# ---------------------------------------------------------------------------

def _download_template_from_supabase() -> bytes:
    """Download best available template (PDF first, DOCX fallback)."""
    sb = get_supabase_client(prefer_service_role=True)
    candidates = [TEMPLATE_PDF_FILE_NAME, TEMPLATE_FILE_NAME]
    last_err: Exception | None = None
    for fname in candidates:
        try:
            response = sb.storage.from_(SUPABASE_TEMPLATE_BUCKET).download(fname)
            if response and len(response) > 0:
                return bytes(response)
        except Exception as e:
            last_err = e
    raise RuntimeError(
        f"Could not download PTW template from Supabase Storage.\n"
        f"Bucket: {SUPABASE_TEMPLATE_BUCKET}\n"
        f"Tried: {candidates}\n"
        f"Last error: {last_err}\n\n"
        "Ensure at least one template file is uploaded to Supabase Storage."
    )


def _download_pdf_template_from_supabase() -> bytes:
    """Download PDF template only."""
    sb = get_supabase_client(prefer_service_role=True)
    try:
        response = sb.storage.from_(SUPABASE_TEMPLATE_BUCKET).download(
            TEMPLATE_PDF_FILE_NAME
        )
        if response and len(response) > 0:
            data = bytes(response)
            if data[:4] == b"%PDF":
                return data
            raise RuntimeError(f"Downloaded file is not a valid PDF: {TEMPLATE_PDF_FILE_NAME}")
        raise RuntimeError(f"Empty PDF template received: {TEMPLATE_PDF_FILE_NAME}")
    except Exception as e:
        raise RuntimeError(
            f"Bucket: {SUPABASE_TEMPLATE_BUCKET}\n"
            f"File: {TEMPLATE_PDF_FILE_NAME}\n"
            f"Please upload '{TEMPLATE_PDF_FILE_NAME}' to the bucket.\n"
            f"Error: {e}"
        ) from e


# ---------------------------------------------------------------------------
# generate_ptw_docx_from_template (verbatim from S1.py)
# ---------------------------------------------------------------------------

def generate_ptw_docx_from_template(template_bytes: bytes, form_data: dict) -> bytes:
    """
    Fill Electrical PTW Word template with form data.
    User-input values appear in GREEN BOLD.
    Logic verbatim from S1.generate_ptw_docx_from_template.
    """
    from docx import Document
    from docx.shared import RGBColor

    mapping: dict[str, str] = {}
    for k, v in (form_data or {}).items():
        mapping[str(k)] = "" if v is None else str(v)

    doc = Document(BytesIO(template_bytes))
    GREEN_COLOR = RGBColor(0, 128, 0)
    placeholder_pattern = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")

    def get_full_paragraph_text(paragraph) -> str:
        runs = getattr(paragraph, "runs", None)
        if runs:
            return "".join(r.text or "" for r in runs)
        return getattr(paragraph, "text", "") or ""

    def replace_in_paragraph_with_formatting(paragraph) -> None:
        full_text = get_full_paragraph_text(paragraph)
        if not full_text:
            full_text = getattr(paragraph, "text", "") or ""
        if not full_text:
            return
        matches = list(placeholder_pattern.finditer(full_text))
        if not matches:
            return
        parts = []
        last_end = 0
        for m in matches:
            if m.start() > last_end:
                parts.append(("static", full_text[last_end : m.start()]))
            key = m.group(1).strip()
            replacement = mapping.get(key, "")
            if replacement:
                parts.append(("user", replacement))
            last_end = m.end()
        if last_end < len(full_text):
            parts.append(("static", full_text[last_end:]))
        p_element = paragraph._element
        for child in list(p_element):
            tag_name = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag_name == "r":
                p_element.remove(child)
        for part_type, part_text in parts:
            if not part_text:
                continue
            new_run = paragraph.add_run(part_text)
            if part_type == "user":
                new_run.bold = True
                new_run.font.color.rgb = GREEN_COLOR

    def replace_in_cell(cell) -> None:
        for p in cell.paragraphs:
            replace_in_paragraph_with_formatting(p)
        for t in cell.tables:
            replace_in_table(t)

    def replace_in_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell)

    for p in doc.paragraphs:
        replace_in_paragraph_with_formatting(p)
    for t in doc.tables:
        replace_in_table(t)
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph_with_formatting(p)
        for t in section.header.tables:
            replace_in_table(t)
        for p in section.footer.paragraphs:
            replace_in_paragraph_with_formatting(p)
        for t in section.footer.tables:
            replace_in_table(t)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ---------------------------------------------------------------------------
# convert_docx_to_pdf (verbatim from S1.py, LibreOffice path)
# ---------------------------------------------------------------------------

def convert_docx_to_pdf(
    docx_bytes: bytes, *, progress_callback: Optional[Callable] = None
) -> bytes:
    """
    Convert DOCX bytes → PDF bytes.
    Method priority: LibreOffice (Docker has it installed) → return DOCX as fallback.
    Verbatim logic from S1.convert_docx_to_pdf (Linux paths only in FastAPI context).
    """
    import shutil

    temp_dir = Path(tempfile.mkdtemp())
    docx_path = temp_dir / f"ptw_temp_{os.getpid()}.docx"
    pdf_path = temp_dir / f"ptw_temp_{os.getpid()}.pdf"

    try:
        docx_path.write_bytes(docx_bytes)

        soffice_paths: list[str] = []
        for cmd in ("soffice", "libreoffice"):
            p = shutil.which(cmd)
            if p:
                soffice_paths.append(p)
        soffice_paths.extend(
            [
                "/usr/bin/soffice",
                "/usr/bin/libreoffice",
                "/snap/bin/libreoffice",
                "/usr/lib/libreoffice/program/soffice",
            ]
        )
        seen: set[str] = set()
        soffice_paths = [
            p for p in soffice_paths if p and (p not in seen and not seen.add(p))  # type: ignore[func-returns-value]
        ]

        for sp in soffice_paths:
            if not Path(sp).exists():
                continue
            try:
                subprocess.run(
                    [
                        sp,
                        "--headless",
                        "--nologo",
                        "--nolockcheck",
                        "--norestore",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(temp_dir),
                        str(docx_path),
                    ],
                    capture_output=True,
                    timeout=120,
                )
                if pdf_path.exists() and pdf_path.stat().st_size > 0:
                    return pdf_path.read_bytes()
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed ({sp}): {e}")

    finally:
        try:
            if docx_path.exists():
                docx_path.unlink()
            if pdf_path.exists():
                pdf_path.unlink()
            temp_dir.rmdir()
        except Exception:
            pass

    return docx_bytes  # fallback: caller will raise if not PDF


# ---------------------------------------------------------------------------
# generate_ptw_pdf (verbatim dispatch logic from S1.py)
# ---------------------------------------------------------------------------

def generate_ptw_pdf(
    template_bytes: bytes,
    form_data: dict,
    *,
    progress_callback: Optional[Callable] = None,
) -> bytes:
    """
    Generate PTW PDF.
    Primary: PDF template + ReportLab overlay (no system binaries).
    Fallback: DOCX → LibreOffice → PDF.
    Dispatch logic verbatim from S1.generate_ptw_pdf.
    """
    # Primary: ReportLab overlay
    try:
        if progress_callback:
            try:
                progress_callback(5, "Generating PDF from template...")
            except Exception:
                pass
        from app.services.permit.s1_overlay import generate_ptw_pdf_from_template

        pdf_bytes = generate_ptw_pdf_from_template(
            form_data, progress_callback=progress_callback
        )
        if isinstance(pdf_bytes, (bytes, bytearray)) and pdf_bytes[:4] == b"%PDF":
            return bytes(pdf_bytes)
    except Exception as template_err:
        logger.warning(f"PDF template overlay failed: {template_err}")

    # Fallback: DOCX conversion
    try:
        docx_bytes = generate_ptw_docx_from_template(template_bytes, form_data)
        if progress_callback:
            try:
                progress_callback(60, "Converting DOCX → PDF...")
            except Exception:
                pass
        pdf_bytes = convert_docx_to_pdf(docx_bytes, progress_callback=progress_callback)
        if isinstance(pdf_bytes, (bytes, bytearray)) and pdf_bytes[:4] == b"%PDF":
            return bytes(pdf_bytes)
    except Exception as docx_err:
        logger.warning(f"DOCX conversion fallback failed: {docx_err}")

    raise RuntimeError(
        "PDF generation failed.\n\n"
        "Please ensure the PDF template is uploaded to Supabase Storage:\n"
        f"  Bucket: {SUPABASE_TEMPLATE_BUCKET}\n"
        f"  File: {TEMPLATE_PDF_FILE_NAME}\n\n"
        "The PDF template is required for hosted deployments."
    )
